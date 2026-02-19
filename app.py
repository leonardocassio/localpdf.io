import io
import os
import shutil
import tempfile
import threading
import uuid
import zipfile

import fitz  # PyMuPDF
import ghostscript
import openpyxl
from flask import Flask, jsonify, render_template_string, request, send_file
from pdf2docx import Converter
from pdf2docx.converter import ConversionException
from PIL import Image
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 100 * 1024 * 1024  # 100MB max
app.config["UPLOAD_FOLDER"] = "uploads"
app.config["OUTPUT_FOLDER"] = "outputs"

os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)
os.makedirs(app.config["OUTPUT_FOLDER"], exist_ok=True)

ALLOWED_EXTENSIONS = {"pdf", "docx", "txt", "xlsx", "jpg", "jpeg", "png"}

# Gerenciamento de tarefas em background
tasks: dict = {}
tasks_lock = threading.Lock()


def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


class SavedFile:
    def __init__(self, path: str):
        self.path = path
        self.filename = os.path.basename(path)

    def save(self, dest: str):
        if os.path.abspath(dest) != os.path.abspath(self.path):
            shutil.copy(self.path, dest)


def set_progress(task_id: str, progress: int, message: str = "", status: str = "processing"):
    with tasks_lock:
        if task_id in tasks:
            tasks[task_id]["progress"] = progress
            tasks[task_id]["message"]  = message
            tasks[task_id]["status"]   = status

HTML_TEMPLATE = """

<!DOCTYPE html>
<html lang="pt-BR">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>NeoConvert - Ferramentas PDF Corporativas</title>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&display=swap" rel="stylesheet">
    <style>
        * { 
            margin: 0; 
            padding: 0; 
            box-sizing: border-box; 
        }
        
        :root {
            --primary: #0066CC;
            --primary-dark: #004C99;
            --primary-light: #3385D6;
            --secondary: #00A896;
            --secondary-dark: #008778;
            --accent: #0099FF;
            --dark: #1A1A2E;
            --gray: #64748B;
            --light-gray: #F1F5F9;
            --white: #FFFFFF;
            --success: #10B981;
            --error: #EF4444;
            --warning: #F59E0B;
        }
        
        body { 
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; 
            background: linear-gradient(135deg, #0066CC 0%, #00A896 100%);
            min-height: 100vh; 
            position: relative;
            overflow-x: hidden;
        }
        
        /* Animated background */
        body::before {
            content: '';
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            height: 100%;
            background: 
                radial-gradient(circle at 20% 50%, rgba(0, 168, 150, 0.3) 0%, transparent 50%),
                radial-gradient(circle at 80% 80%, rgba(0, 102, 204, 0.3) 0%, transparent 50%),
                radial-gradient(circle at 40% 20%, rgba(0, 153, 255, 0.2) 0%, transparent 50%);
            animation: gradientShift 15s ease infinite;
            pointer-events: none;
            z-index: 0;
        }
        
        @keyframes gradientShift {
            0%, 100% { opacity: 1; transform: scale(1); }
            50% { opacity: 0.8; transform: scale(1.1); }
        }
        
        .container { 
            max-width: 1400px; 
            margin: 0 auto; 
            padding: 20px; 
            position: relative;
            z-index: 1;
        }
        
        /* Header Styles */
        .header { 
            text-align: center; 
            color: white; 
            margin-bottom: 50px; 
            padding: 60px 20px 40px;
            animation: fadeInDown 0.8s ease;
        }
        
        @keyframes fadeInDown {
            from { opacity: 0; transform: translateY(-30px); }
            to { opacity: 1; transform: translateY(0); }
        }
        
        .header .badge {
            display: inline-block;
            background: rgba(255, 255, 255, 0.2);
            backdrop-filter: blur(20px);
            -webkit-backdrop-filter: blur(20px);
            padding: 10px 24px;
            border-radius: 50px;
            margin-bottom: 24px;
            font-size: 0.9rem;
            font-weight: 600;
            border: 1px solid rgba(255, 255, 255, 0.3);
            box-shadow: 0 8px 32px rgba(0, 0, 0, 0.1);
            animation: pulse 2s ease-in-out infinite;
        }
        
        @keyframes pulse {
            0%, 100% { transform: scale(1); }
            50% { transform: scale(1.05); }
        }
        
        .header h1 { 
            font-size: 4em; 
            margin-bottom: 16px; 
            font-weight: 900;
            letter-spacing: -2px;
            text-shadow: 0 4px 20px rgba(0, 0, 0, 0.2);
            background: linear-gradient(to right, #FFFFFF, #E0F2FE);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
        }
        
        .header p { 
            font-size: 1.3em; 
            opacity: 0.95; 
            font-weight: 400;
            text-shadow: 0 2px 10px rgba(0, 0, 0, 0.2);
            max-width: 600px;
            margin: 0 auto;
            line-height: 1.6;
        }
        
        /* Tools Grid */
        .tools-grid { 
            display: grid; 
            grid-template-columns: repeat(auto-fit, minmax(320px, 1fr)); 
            gap: 24px; 
            margin-bottom: 50px;
            animation: fadeInUp 0.8s ease;
        }
        
        @keyframes fadeInUp {
            from { opacity: 0; transform: translateY(30px); }
            to { opacity: 1; transform: translateY(0); }
        }
        
        /* Tool Card */
        .tool-card { 
            background: rgba(255, 255, 255, 0.95);
            backdrop-filter: blur(20px);
            -webkit-backdrop-filter: blur(20px);
            border-radius: 20px; 
            padding: 36px; 
            text-align: center; 
            box-shadow: 
                0 10px 40px rgba(0, 0, 0, 0.1),
                0 2px 8px rgba(0, 0, 0, 0.06);
            transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1); 
            cursor: pointer; 
            border: 1px solid rgba(255, 255, 255, 0.8);
            position: relative;
            overflow: hidden;
        }
        
        .tool-card::before {
            content: '';
            position: absolute;
            top: 0;
            left: -100%;
            width: 100%;
            height: 100%;
            background: linear-gradient(90deg, transparent, rgba(255,255,255,0.3), transparent);
            transition: left 0.5s;
        }
        
        .tool-card:hover::before {
            left: 100%;
        }
        
        .tool-card:hover { 
            transform: translateY(-12px) scale(1.02); 
            box-shadow: 
                0 20px 60px rgba(0, 102, 204, 0.3),
                0 8px 16px rgba(0, 0, 0, 0.1);
            border-color: var(--primary);
        }
        
        .tool-card .icon {
            font-size: 3.5em;
            margin-bottom: 20px;
            display: inline-block;
            transition: transform 0.3s ease;
            filter: drop-shadow(0 4px 8px rgba(0, 0, 0, 0.1));
        }
        
        .tool-card:hover .icon {
            transform: scale(1.1) rotate(5deg);
        }
        
        .tool-card h3 { 
            color: var(--primary); 
            margin-bottom: 16px; 
            font-size: 1.5em; 
            font-weight: 700;
            letter-spacing: -0.5px;
        }
        
        .tool-card p { 
            color: var(--gray); 
            line-height: 1.7;
            font-size: 1.05em;
            font-weight: 400;
        }
        
        /* Upload Area */
        .upload-area { 
            border: 3px dashed #CBD5E1; 
            border-radius: 20px; 
            padding: 60px 40px; 
            text-align: center; 
            background: linear-gradient(135deg, #F8FAFC 0%, #F1F5F9 100%);
            margin: 30px 0; 
            transition: all 0.3s ease; 
            cursor: pointer;
            position: relative;
        }
        
        .upload-area::before {
            content: '📤';
            position: absolute;
            top: 50%;
            left: 50%;
            transform: translate(-50%, -50%);
            font-size: 8em;
            opacity: 0.05;
            pointer-events: none;
        }
        
        .upload-area:hover { 
            border-color: var(--primary); 
            background: linear-gradient(135deg, #EFF6FF 0%, #DBEAFE 100%);
            border-width: 3px;
            transform: scale(1.01);
        }
        
        .upload-area.dragover { 
            border-color: var(--secondary); 
            background: linear-gradient(135deg, #ECFDF5 0%, #D1FAE5 100%); 
            border-width: 4px;
            transform: scale(1.02);
        }
        
        .upload-area .upload-icon {
            font-size: 4em;
            margin-bottom: 20px;
            display: block;
            animation: bounce 2s ease-in-out infinite;
        }
        
        @keyframes bounce {
            0%, 100% { transform: translateY(0); }
            50% { transform: translateY(-10px); }
        }
        
        .upload-area p {
            color: var(--gray);
            font-size: 1.2em;
            margin-bottom: 20px;
            font-weight: 500;
        }
        
        .file-input { display: none; }
        
        .upload-btn { 
            background: linear-gradient(135deg, var(--primary) 0%, var(--primary-light) 100%);
            color: white; 
            padding: 16px 40px; 
            border: none; 
            border-radius: 12px; 
            cursor: pointer; 
            font-size: 1.1em; 
            font-weight: 700;
            transition: all 0.3s ease; 
            box-shadow: 0 4px 20px rgba(0, 102, 204, 0.3);
            letter-spacing: 0.5px;
        }
        
        .upload-btn:hover { 
            background: linear-gradient(135deg, var(--primary-dark) 0%, var(--primary) 100%);
            transform: translateY(-3px);
            box-shadow: 0 8px 30px rgba(0, 102, 204, 0.4);
        }
        
        .upload-btn:active {
            transform: translateY(-1px);
        }
        
        .convert-btn { 
            background: linear-gradient(135deg, var(--secondary) 0%, var(--secondary-dark) 100%);
            color: white; 
            padding: 18px 50px; 
            border: none; 
            border-radius: 12px; 
            cursor: pointer; 
            font-size: 1.3em; 
            font-weight: 700;
            margin-top: 30px; 
            transition: all 0.3s ease; 
            box-shadow: 0 6px 25px rgba(0, 168, 150, 0.3);
            letter-spacing: 0.5px;
        }
        
        .convert-btn:hover { 
            background: linear-gradient(135deg, var(--secondary-dark) 0%, #006D5F 100%);
            transform: translateY(-4px);
            box-shadow: 0 10px 35px rgba(0, 168, 150, 0.4);
        }
        
        .convert-btn:disabled { 
            background: linear-gradient(135deg, #CBD5E1 0%, #94A3B8 100%);
            cursor: not-allowed; 
            transform: none;
            box-shadow: none;
        }
        
        /* File List */
        .file-list { 
            margin-top: 30px; 
        }
        
        .file-item { 
            background: white;
            padding: 18px 24px; 
            margin: 12px 0; 
            border-radius: 12px; 
            display: flex; 
            justify-content: space-between; 
            align-items: center; 
            border: 2px solid #E2E8F0;
            transition: all 0.3s ease;
            animation: slideIn 0.3s ease;
        }
        
        @keyframes slideIn {
            from { opacity: 0; transform: translateX(-20px); }
            to { opacity: 1; transform: translateX(0); }
        }
        
        .file-item:hover {
            border-color: var(--primary);
            box-shadow: 0 4px 15px rgba(0, 102, 204, 0.1);
            transform: translateX(5px);
        }
        
        .file-item span {
            color: var(--dark);
            font-weight: 600;
            display: flex;
            align-items: center;
            gap: 10px;
        }
        
        .file-item span::before {
            content: '📄';
            font-size: 1.5em;
        }
        
        /* Progress Bar */
        .progress { 
            width: 100%; 
            background: #E2E8F0; 
            border-radius: 50px; 
            margin: 30px 0; 
            height: 30px;
            overflow: hidden;
            box-shadow: inset 0 2px 8px rgba(0, 0, 0, 0.1);
        }
        
        .progress-bar { 
            height: 100%; 
            background: linear-gradient(90deg, var(--primary) 0%, var(--secondary) 100%);
            border-radius: 50px; 
            width: 0%; 
            transition: width 0.3s ease;
            position: relative;
            overflow: hidden;
            animation: progressAnimation 1.5s ease infinite;
        }
        
        @keyframes progressAnimation {
            0% { background-position: 0% 50%; }
            50% { background-position: 100% 50%; }
            100% { background-position: 0% 50%; }
        }
        
        .progress-bar::before {
            content: '';
            position: absolute;
            top: 0;
            left: -100%;
            width: 100%;
            height: 100%;
            background: linear-gradient(90deg, transparent, rgba(255,255,255,0.3), transparent);
            animation: shimmer 2s infinite;
        }
        
        @keyframes shimmer {
            to { left: 100%; }
        }
        
        /* Result Messages */
        .result { 
            margin-top: 30px; 
            padding: 24px 28px; 
            background: linear-gradient(135deg, #D1FAE5 0%, #A7F3D0 100%);
            border-radius: 16px; 
            color: #065F46; 
            border: 2px solid var(--success);
            animation: slideIn 0.4s ease;
            box-shadow: 0 4px 20px rgba(16, 185, 129, 0.2);
        }
        
        .result h4 {
            margin-bottom: 12px;
            font-size: 1.3em;
            font-weight: 700;
            display: flex;
            align-items: center;
            gap: 10px;
        }
        
        .result p {
            font-size: 1.05em;
            line-height: 1.6;
        }
        
        .error { 
            margin-top: 30px; 
            padding: 24px 28px; 
            background: linear-gradient(135deg, #FEE2E2 0%, #FECACA 100%);
            border-radius: 16px; 
            color: #991B1B; 
            border: 2px solid var(--error);
            animation: shake 0.5s ease;
            box-shadow: 0 4px 20px rgba(239, 68, 68, 0.2);
        }
        
        @keyframes shake {
            0%, 100% { transform: translateX(0); }
            25% { transform: translateX(-10px); }
            75% { transform: translateX(10px); }
        }
        
        .error h4 {
            margin-bottom: 12px;
            font-size: 1.3em;
            font-weight: 700;
        }
        
        .hidden { display: none; }
        
        /* Back Button */
        .back-btn { 
            background: rgba(255, 255, 255, 0.95);
            color: var(--primary); 
            padding: 12px 28px; 
            border: 2px solid var(--primary);
            border-radius: 12px; 
            cursor: pointer; 
            margin-bottom: 30px; 
            font-weight: 700;
            font-size: 1.05em;
            transition: all 0.3s ease;
            display: inline-flex;
            align-items: center;
            gap: 8px;
            box-shadow: 0 4px 15px rgba(0, 102, 204, 0.2);
        }
        
        .back-btn:hover { 
            background: var(--primary);
            color: white;
            transform: translateX(-5px);
            box-shadow: 0 6px 20px rgba(0, 102, 204, 0.3);
        }
        
        /* Remove Button */
        .remove-btn {
            background: linear-gradient(135deg, var(--error) 0%, #DC2626 100%);
            color: white;
            border: none;
            padding: 8px 20px;
            border-radius: 8px;
            cursor: pointer;
            font-weight: 700;
            font-size: 0.95em;
            transition: all 0.3s ease;
            box-shadow: 0 4px 15px rgba(239, 68, 68, 0.2);
        }
        
        .remove-btn:hover {
            background: linear-gradient(135deg, #DC2626 0%, #B91C1C 100%);
            transform: translateY(-2px);
            box-shadow: 0 6px 20px rgba(239, 68, 68, 0.3);
        }
        
        /* Footer */
        .footer { 
            text-align: center; 
            color: white; 
            margin-top: 80px; 
            padding: 40px 20px; 
            border-top: 1px solid rgba(255, 255, 255, 0.2);
            background: rgba(0, 0, 0, 0.1);
            backdrop-filter: blur(20px);
            -webkit-backdrop-filter: blur(20px);
            border-radius: 20px 20px 0 0;
        }
        
        .footer p { 
            margin-bottom: 12px; 
            opacity: 0.95;
            font-size: 1.05em;
            font-weight: 500;
        }
        
        .footer a { 
            color: #FFFFFF; 
            text-decoration: none; 
            font-weight: 700;
            transition: all 0.3s;
            padding: 4px 8px;
            border-radius: 6px;
        }
        
        .footer a:hover { 
            background: rgba(255, 255, 255, 0.2);
            transform: translateY(-2px);
        }
        
        .footer .security-note {
            margin-top: 20px;
            padding: 16px 24px;
            background: rgba(255, 255, 255, 0.15);
            backdrop-filter: blur(10px);
            border-radius: 12px;
            display: inline-block;
            font-size: 0.95em;
            border: 1px solid rgba(255, 255, 255, 0.2);
        }
        
        /* Responsive */
        @media (max-width: 768px) {
            .header h1 { font-size: 2.5em; }
            .tools-grid { grid-template-columns: 1fr; }
            .tool-card { padding: 28px; }
            .upload-area { padding: 40px 20px; }
            .container { padding: 15px; }
        }
        
        /* Loading Animation */
        @keyframes spin {
            to { transform: rotate(360deg); }
        }
        
        .loading {
            display: inline-block;
            width: 20px;
            height: 20px;
            border: 3px solid rgba(255, 255, 255, 0.3);
            border-radius: 50%;
            border-top-color: white;
            animation: spin 0.8s linear infinite;
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <div class="badge">🔒 100% Local & Seguro</div>
            <h1>📄 NeoConvert</h1>
            <p>Ferramentas PDF corporativas com total privacidade e segurança</p>
        </div>

        <div id="home-view">
            <div class="tools-grid">
                <div class="tool-card" onclick="showTool('pdf-to-images')">
                    <div class="icon">🖼️</div>
                    <h3>PDF para Imagens</h3>
                    <p>Converta páginas PDF em imagens JPG ou PNG</p>
                </div>
                <div class="tool-card" onclick="showTool('images-to-pdf')">
                    <div class="icon">📄</div>
                    <h3>Imagens para PDF</h3>
                    <p>Combine várias imagens em um único PDF</p>
                </div>
                <div class="tool-card" onclick="showTool('merge-pdf')">
                    <div class="icon">🔗</div>
                    <h3>Mesclar PDFs</h3>
                    <p>Combine vários PDFs em um documento único</p>
                </div>
                <div class="tool-card" onclick="showTool('split-pdf')">
                    <div class="icon">✂️</div>
                    <h3>Dividir PDF</h3>
                    <p>Extraia páginas específicas do seu PDF</p>
                </div>
                <div class="tool-card" onclick="showTool('compress-pdf')">
                    <div class="icon">📦</div>
                    <h3>Comprimir PDF</h3>
                    <p>Reduza o tamanho do seu arquivo PDF</p>
                </div>
                <div class="tool-card" onclick="showTool('pdf-to-pdfa')">
                    <div class="icon">🔒</div>
                    <h3>PDF para PDF/A</h3>
                    <p>Padronize seu PDF para arquivamento</p>
                </div>
                <div class="tool-card" onclick="showTool('word-to-pdf')">
                    <div class="icon">📝</div>
                    <h3>Word para PDF</h3>
                    <p>Converta documentos DOCX para PDF</p>
                </div>
                <div class="tool-card" onclick="showTool('excel-to-pdf')">
                    <div class="icon">📊</div>
                    <h3>Excel para PDF</h3>
                    <p>Converta planilhas XLSX para PDF</p>
                </div>
                <div class="tool-card" onclick="showTool('txt-to-pdf')">
                    <div class="icon">📃</div>
                    <h3>TXT para PDF</h3>
                    <p>Converta arquivos de texto para PDF</p>
                </div>
                <div class="tool-card" onclick="showTool('pdf-to-word')">
                    <div class="icon">🔄</div>
                    <h3>PDF para Word</h3>
                    <p>Converta PDF para Word editável</p>
                </div>
            </div>
        </div>

        <!-- Tool Views -->
        <div id="tool-views" class="hidden">
            <button class="back-btn" onclick="showHome()">&larr; Voltar</button>
            <div class="tool-card">
                <h3 id="tool-title"></h3>
                <p id="tool-description"></p>

                <div class="upload-area" id="upload-area" onclick="document.getElementById('file-input').click()">
                    <input type="file" id="file-input" class="file-input" multiple accept=".pdf,.docx,.jpg,.jpeg,.png,.txt,.xlsx">
                    <span class="upload-icon">📁</span>
                    <p>Clique aqui ou arraste arquivos para fazer upload</p>
                    <button class="upload-btn">Escolher Arquivos</button>
                </div>

                <div id="file-list" class="file-list"></div>

                <button id="convert-btn" class="convert-btn hidden" onclick="convertFiles()">🚀 Converter Agora</button>

                <div id="progress" class="progress hidden">
                    <div id="progress-bar" class="progress-bar"></div>
                </div>
                <p id="progress-msg" style="text-align:center;color:var(--gray);margin:8px 0 0;font-size:0.9em;min-height:1.3em;"></p>

                <div id="result" class="hidden"></div>
            </div>
        </div>

        <div class="footer">
            <p><strong>NeoConvert</strong> - Ferramenta Corporativa Interna</p>
            <p>
                <a href="mailto:ti-infra@neogenomica.com.br">✉️ ti-infra@neogenomica.com.br</a>
            </p>
            <div class="security-note">
                🛡️ Processamento 100% local • Seus arquivos nunca saem da infraestrutura interna
            </div>
        </div>
    </div>

    <script>
        let currentTool = '';
        let uploadedFiles = [];

        const tools = {
            'pdf-to-images': { title: 'PDF para Imagens',    description: 'Converta cada pagina do seu PDF em imagens separadas', accept: '.pdf',              multiple: false },
            'images-to-pdf': { title: 'Imagens para PDF',    description: 'Combine multiplas imagens em um unico arquivo PDF',    accept: '.jpg,.jpeg,.png',   multiple: true  },
            'merge-pdf':     { title: 'Mesclar PDFs',        description: 'Combine varios arquivos PDF em um documento unico',    accept: '.pdf',              multiple: true  },
            'split-pdf':     { title: 'Dividir PDF',         description: 'Extraia paginas especificas do seu PDF',               accept: '.pdf',              multiple: false },
            'compress-pdf':  { title: 'Comprimir PDF',       description: 'Reduza o tamanho do arquivo PDF mantendo a qualidade', accept: '.pdf',              multiple: false },
            'pdf-to-pdfa':   { title: 'PDF para PDF/A',      description: 'Converta PDFs para o padrao de arquivamento PDF/A-1b', accept: '.pdf',              multiple: true  },
            'word-to-pdf':   { title: 'Word para PDF',       description: 'Converta documentos Word (.docx) para PDF',           accept: '.docx',             multiple: true  },
            'excel-to-pdf':  { title: 'Excel para PDF',      description: 'Converta planilhas Excel (.xlsx) para PDF',           accept: '.xlsx',             multiple: false },
            'txt-to-pdf':    { title: 'TXT para PDF',        description: 'Converta arquivos de texto simples (.txt) para PDF',  accept: '.txt',              multiple: false },
            'pdf-to-word':   { title: 'PDF para Word',       description: 'Converta seus documentos PDF para Word (.docx)',      accept: '.pdf',              multiple: false }
        };

        function showTool(toolName) {
            currentTool = toolName;
            const tool = tools[toolName];
            document.getElementById('home-view').classList.add('hidden');
            document.getElementById('tool-views').classList.remove('hidden');
            document.getElementById('tool-title').innerText       = tool.title;
            document.getElementById('tool-description').innerText = tool.description;
            document.getElementById('file-input').accept   = tool.accept;
            document.getElementById('file-input').multiple = tool.multiple;
            uploadedFiles = [];
            updateFileList();
            hideResult();
        }

        function showHome() {
            document.getElementById('home-view').classList.remove('hidden');
            document.getElementById('tool-views').classList.add('hidden');
            uploadedFiles = [];
        }

        function updateFileList() {
            const fileList   = document.getElementById('file-list');
            const convertBtn = document.getElementById('convert-btn');
            if (uploadedFiles.length === 0) {
                fileList.innerHTML = '';
                convertBtn.classList.add('hidden');
                return;
            }
            fileList.innerHTML = uploadedFiles.map((file, index) => `
                <div class="file-item">
                    <span>${file.name} <small style="opacity:0.7">(${(file.size/1024/1024).toFixed(2)} MB)</small></span>
                    <button onclick="removeFile(${index})" class="remove-btn">Remover</button>
                </div>`).join('');
            convertBtn.classList.remove('hidden');
        }

        function removeFile(index) {
            uploadedFiles.splice(index, 1);
            updateFileList();
        }

        function hideResult() {
            document.getElementById('result').classList.add('hidden');
            document.getElementById('progress').classList.add('hidden');
            const msg = document.getElementById('progress-msg');
            if (msg) msg.textContent = '';
        }

        document.getElementById('file-input').addEventListener('change', function(e) {
            const files = Array.from(e.target.files);
            uploadedFiles = tools[currentTool].multiple ? uploadedFiles.concat(files) : files.slice(0,1);
            updateFileList();
        });

        const uploadArea = document.getElementById('upload-area');
        uploadArea.addEventListener('dragover',  e => { e.preventDefault(); uploadArea.classList.add('dragover'); });
        uploadArea.addEventListener('dragleave', e => { e.preventDefault(); uploadArea.classList.remove('dragover'); });
        uploadArea.addEventListener('drop', function(e) {
            e.preventDefault();
            uploadArea.classList.remove('dragover');
            const files = Array.from(e.dataTransfer.files);
            uploadedFiles = tools[currentTool].multiple ? uploadedFiles.concat(files) : files.slice(0,1);
            updateFileList();
        });

        async function convertFiles() {
            if (uploadedFiles.length === 0) return;

            const formData = new FormData();
            uploadedFiles.forEach(file => formData.append('files', file));
            formData.append('tool', currentTool);

            const progressBar = document.getElementById('progress-bar');
            const progressMsg = document.getElementById('progress-msg');
            document.getElementById('progress').classList.remove('hidden');
            document.getElementById('convert-btn').disabled = true;
            hideResult();

            progressBar.style.width = '5%';
            if (progressMsg) progressMsg.textContent = 'Enviando arquivos...';

            try {
                const response = await fetch('/convert', { method: 'POST', body: formData });
                if (!response.ok) {
                    const err = await response.json();
                    throw new Error(err.error || 'Erro ao iniciar conversao');
                }

                const { task_id } = await response.json();

                await new Promise((resolve, reject) => {
                    const interval = setInterval(async () => {
                        try {
                            const res = await fetch(`/progress/${task_id}`);
                            if (!res.ok) { clearInterval(interval); reject(new Error('Erro ao verificar progresso')); return; }
                            const prog = await res.json();
                            progressBar.style.width = prog.progress + '%';
                            if (progressMsg) progressMsg.textContent = prog.message;
                            if (prog.status === 'done')  { clearInterval(interval); progressBar.style.width = '100%'; resolve(); }
                            if (prog.status === 'error') { clearInterval(interval); reject(new Error(prog.message)); }
                        } catch(e) { clearInterval(interval); reject(e); }
                    }, 600);
                });

                if (progressMsg) progressMsg.textContent = 'Baixando arquivo...';
                const a = document.createElement('a');
                a.href = `/download/${task_id}`;
                a.download = '';
                document.body.appendChild(a); a.click(); document.body.removeChild(a);

                document.getElementById('result').className = 'result';
                document.getElementById('result').innerHTML  = '<h4>Sucesso!</h4><p>Arquivo convertido e baixado com sucesso!</p>';
                document.getElementById('result').classList.remove('hidden');

            } catch (error) {
                document.getElementById('result').className = 'error';
                document.getElementById('result').innerHTML  = `<h4>Erro!</h4><p>${error.message || 'Ocorreu um erro. Tente novamente.'}</p>`;
                document.getElementById('result').classList.remove('hidden');
            } finally {
                setTimeout(() => {
                    document.getElementById('progress').classList.add('hidden');
                    progressBar.style.width = '0%';
                    if (progressMsg) progressMsg.textContent = '';
                }, 2000);
                document.getElementById('convert-btn').disabled = false;
            }
        }
    </script></html>
"""


# [Resto do código Python permanece igual - funções de conversão]

@app.route("/")
def index():
    return render_template_string(HTML_TEMPLATE)


def excel_to_pdf(file, temp_dir):
    xlsx_path = os.path.join(temp_dir, secure_filename(file.filename))
    file.save(xlsx_path)

    pdf_path = os.path.join(temp_dir, "excel_to_pdf.pdf")
    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    y_position = height - 50

    try:
        workbook = openpyxl.load_workbook(xlsx_path)
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            c.setFont("Helvetica", 10)
            c.drawString(50, y_position, f"--- Planilha: {sheet_name} ---")
            y_position -= 20

            for row_idx, row in enumerate(sheet.iter_rows()):
                row_data = [
                    str(cell.value) if cell.value is not None else "" for cell in row
                ]
                line_text = " | ".join(row_data)

                max_line_width = int((width - 100) / 6)
                if len(line_text) > max_line_width:
                    line_text = line_text[:max_line_width] + "..."

                if y_position < 50:
                    c.showPage()
                    y_position = height - 50
                    c.setFont("Helvetica", 10)

                c.drawString(50, y_position, line_text)
                y_position -= 15

            y_position -= 30
            if y_position < 50 and sheet_name != workbook.sheetnames[-1]:
                c.showPage()
                y_position = height - 50

    except Exception as e:
        c.drawString(50, y_position - 20, f"Erro ao ler planilha: {e}")
        print(f"Erro ao ler planilha Excel: {e}")

    c.save()
    return [pdf_path]


def txt_to_pdf(file, temp_dir):
    txt_path = os.path.join(temp_dir, secure_filename(file.filename))
    file.save(txt_path)

    pdf_path = os.path.join(temp_dir, "text_to_pdf.pdf")
    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    y_position = height - 50

    c.setFont("Helvetica", 12)

    try:
        with open(txt_path, "r", encoding="utf-8") as f:
            for line in f:
                text_line = line.strip()
                max_width_px = width - 100

                approx_char_width_px = 7
                chars_per_line = int(max_width_px / approx_char_width_px)

                if len(text_line) > chars_per_line:
                    chunks = [
                        text_line[i : i + chars_per_line]
                        for i in range(0, len(text_line), chars_per_line)
                    ]
                else:
                    chunks = [text_line]

                for chunk in chunks:
                    if y_position < 50:
                        c.showPage()
                        y_position = height - 50
                        c.setFont("Helvetica", 12)

                    c.drawString(50, y_position, chunk)
                    y_position -= 15

    except Exception as e:
        c.drawString(50, y_position - 20, f"Erro ao ler arquivo de texto: {e}")
        print(f"Erro ao ler arquivo de texto: {e}")

    c.save()
    return [pdf_path]


@app.route("/convert", methods=["POST"])
def convert():
    if "files" not in request.files:
        return jsonify({"error": "Nenhum arquivo enviado"}), 400

    files = request.files.getlist("files")
    tool = request.form.get("tool")

    if not files or files[0].filename == "":
        return jsonify({"error": "Nenhum arquivo selecionado"}), 400

    for f in files:
        if not allowed_file(f.filename):
            return jsonify({"error": f"Extensão não permitida: {f.filename}"}), 400

    temp_dir = tempfile.mkdtemp()
    response = None
    try:
        if tool == "pdf-to-images":
            output_files = pdf_to_images(files[0], temp_dir)
        elif tool == "images-to-pdf":
            output_files = images_to_pdf(files, temp_dir)
        elif tool == "merge-pdf":
            output_files = merge_pdfs(files, temp_dir)
        elif tool == "split-pdf":
            output_files = split_pdf(files[0], temp_dir)
        elif tool == "compress-pdf":
            output_files = compress_pdf(files[0], temp_dir)
        elif tool == "pdf-to-pdfa":
            output_files = pdf_to_pdfa(files, temp_dir)
        elif tool == "word-to-pdf":
            output_files = word_to_pdf(files, temp_dir)
        elif tool == "excel-to-pdf":
            output_files = excel_to_pdf(files[0], temp_dir)
        elif tool == "txt-to-pdf":
            output_files = txt_to_pdf(files[0], temp_dir)
        elif tool == "pdf-to-word":
            output_files = pdf_to_word(files[0], temp_dir)
        else:
            return jsonify({"error": "Ferramenta não suportada"}), 400

        response = build_response(output_files, temp_dir)
        return response
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)


def pdf_to_images(file, temp_dir):
    pdf_path = os.path.join(temp_dir, secure_filename(file.filename))
    file.save(pdf_path)

    doc = fitz.open(pdf_path)
    output_files = []

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        img_path = os.path.join(temp_dir, f"page_{page_num + 1}.png")
        pix.save(img_path)
        output_files.append(img_path)

    doc.close()
    return output_files


def images_to_pdf(files, temp_dir):
    images = []
    for file in files:
        img_path = os.path.join(temp_dir, secure_filename(file.filename))
        file.save(img_path)
        img = Image.open(img_path)
        if img.mode != "RGB":
            img = img.convert("RGB")
        images.append(img)

    pdf_path = os.path.join(temp_dir, "images_to_pdf.pdf")
    images[0].save(pdf_path, save_all=True, append_images=images[1:])

    return [pdf_path]


def merge_pdfs(files, temp_dir):
    merged_doc = fitz.open()

    for file in files:
        pdf_path = os.path.join(temp_dir, secure_filename(file.filename))
        file.save(pdf_path)
        doc = fitz.open(pdf_path)
        merged_doc.insert_pdf(doc)
        doc.close()

    output_path = os.path.join(temp_dir, "merged.pdf")
    merged_doc.save(output_path)
    merged_doc.close()

    return [output_path]


def split_pdf(file, temp_dir):
    pdf_path = os.path.join(temp_dir, secure_filename(file.filename))
    file.save(pdf_path)

    doc = fitz.open(pdf_path)
    output_files = []

    for page_num in range(len(doc)):
        new_doc = fitz.open()
        new_doc.insert_pdf(doc, from_page=page_num, to_page=page_num)
        output_path = os.path.join(temp_dir, f"page_{page_num + 1}.pdf")
        new_doc.save(output_path)
        new_doc.close()
        output_files.append(output_path)

    doc.close()
    return output_files


def compress_pdf(file, temp_dir):
    pdf_path = os.path.join(temp_dir, secure_filename(file.filename))
    file.save(pdf_path)

    doc = fitz.open(pdf_path)
    output_path = os.path.join(temp_dir, "compressed.pdf")
    doc.save(output_path, garbage=4, deflate=True, clean=True)
    doc.close()

    return [output_path]


def pdf_to_pdfa(files, temp_dir):
    if not isinstance(files, list):
        files = [files]

    output_files = []

    for file in files:
        input_path = os.path.join(temp_dir, secure_filename(file.filename))
        file.save(input_path)

        base_name, _ = os.path.splitext(os.path.basename(input_path))
        output_path = os.path.join(temp_dir, f"{base_name}_pdfa.pdf")

        gs_args = [
            "gs",
            "-dPDFA=1",
            "-dBATCH",
            "-dNOPAUSE",
            "-dNOOUTERSAVE",
            "-dUseCIEColor",
            "-sProcessColorModel=DeviceRGB",
            "-sDEVICE=pdfwrite",
            "-sColorConversionStrategy=UseDeviceIndependentColor",
            "-dPDFACompatibilityPolicy=1",
            f"-sOutputFile={output_path}",
            input_path,
        ]
        gs_args = [
            arg.encode("utf-8") if isinstance(arg, str) else arg for arg in gs_args
        ]

        try:
            ghostscript.Ghostscript(*gs_args)
        except Exception as e:
            raise RuntimeError(
                f"Erro ao converter {file.filename} para PDF/A: {e}"
            ) from e

        output_files.append(output_path)

    return output_files


def word_to_pdf(files, temp_dir):
    from docx import Document

    pdf_path = os.path.join(temp_dir, "word_to_pdf.pdf")
    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    y_position = height - 50

    if not isinstance(files, list):
        files = [files]

    for file_idx, file in enumerate(files):
        docx_path = os.path.join(temp_dir, secure_filename(file.filename))
        file.save(docx_path)

        doc = Document(docx_path)

        if file_idx > 0:
            c.showPage()
            y_position = height - 50

            c.setFont("Helvetica-Bold", 12)
            c.drawString(50, y_position, f"{'=' * 60}")
            y_position -= 20
            c.drawString(50, y_position, f"Documento: {file.filename}")
            y_position -= 20
            c.drawString(50, y_position, f"{'=' * 60}")
            y_position -= 30
            c.setFont("Helvetica", 11)

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text = paragraph.text
                max_width = width - 100

                approx_char_width = 6
                chars_per_line = int(max_width / approx_char_width)

                words = text.split()
                lines = []
                current_line = []

                for word in words:
                    if len(" ".join(current_line + [word])) <= chars_per_line:
                        current_line.append(word)
                    else:
                        if current_line:
                            lines.append(" ".join(current_line))
                            current_line = [word]
                        else:
                            lines.append(word)

                if current_line:
                    lines.append(" ".join(current_line))

                for line in lines:
                    if y_position < 50:
                        c.showPage()
                        y_position = height - 50

                    c.drawString(50, y_position, line)
                    y_position -= 20

        for table in doc.tables:
            y_position -= 10

            if y_position < 100:
                c.showPage()
                y_position = height - 50

            c.setFont("Helvetica", 9)
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])

                if len(row_text) > 100:
                    row_text = row_text[:97] + "..."

                if y_position < 50:
                    c.showPage()
                    y_position = height - 50

                c.drawString(50, y_position, row_text)
                y_position -= 15

            y_position -= 10
            c.setFont("Helvetica", 11)

    c.save()
    return [pdf_path]


def pdf_to_word(file, temp_dir):
    pdf_path = os.path.join(temp_dir, secure_filename(file.filename))
    file.save(pdf_path)

    docx_filename = os.path.splitext(secure_filename(file.filename))[0] + ".docx"
    docx_path = os.path.join(temp_dir, docx_filename)

    cv = None
    try:
        cv = Converter(pdf_path)
        cv.convert(docx_path)
    except ValueError as e:
        raise RuntimeError(f"Erro no arquivo PDF: {e}") from e
    except ConversionException as e:
        raise RuntimeError(f"Erro interno na conversão: {e}") from e
    except Exception as e:
        raise RuntimeError(f"Erro ao converter {file.filename} para Word: {e}") from e
    finally:
        if cv:
            cv.close()

    return [docx_path]


def build_response(output_files, temp_dir):
    if len(output_files) == 1:
        file_path = output_files[0]
        filename = os.path.basename(file_path)
        with open(file_path, "rb") as f:
            data = f.read()
        return send_file(io.BytesIO(data), as_attachment=True, download_name=filename)
    else:
        zip_path = os.path.join(temp_dir, "converted_files.zip")
        with zipfile.ZipFile(zip_path, "w") as zipf:
            for file_path in output_files:
                zipf.write(file_path, os.path.basename(file_path))
        with open(zip_path, "rb") as f:
            data = f.read()
        return send_file(
            io.BytesIO(data), as_attachment=True, download_name="converted_files.zip"
        )


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)"""


@app.route("/")
def index():
    return render_template_string(HTML_TEMPLATE)


@app.route("/convert", methods=["POST"])
def convert():
    if "files" not in request.files:
        return jsonify({"error": "Nenhum arquivo enviado"}), 400

    files = request.files.getlist("files")
    tool  = request.form.get("tool")

    if not files or files[0].filename == "":
        return jsonify({"error": "Nenhum arquivo selecionado"}), 400

    for f in files:
        if not allowed_file(f.filename):
            return jsonify({"error": f"Extensao nao permitida: {f.filename}"}), 400

    task_id  = str(uuid.uuid4())
    temp_dir = tempfile.mkdtemp()

    saved_paths = []
    for f in files:
        path = os.path.join(temp_dir, secure_filename(f.filename))
        f.save(path)
        saved_paths.append(path)

    with tasks_lock:
        tasks[task_id] = {
            "progress":    0,
            "status":      "processing",
            "message":     "Aguardando inicio...",
            "result_path": None,
            "temp_dir":    temp_dir,
        }

    threading.Thread(
        target=_process_in_background,
        args=(task_id, tool, saved_paths, temp_dir),
        daemon=True,
    ).start()

    return jsonify({"task_id": task_id})


@app.route("/progress/<task_id>")
def get_progress(task_id):
    with tasks_lock:
        task = tasks.get(task_id)
    if not task:
        return jsonify({"error": "Task nao encontrada"}), 404
    return jsonify({
        "progress": task["progress"],
        "status":   task["status"],
        "message":  task["message"],
    })


@app.route("/download/<task_id>")
def download_file(task_id):
    with tasks_lock:
        task = tasks.get(task_id)
    if not task:
        return jsonify({"error": "Task nao encontrada"}), 404
    if task["status"] != "done":
        return jsonify({"error": "Arquivo ainda nao esta pronto"}), 202
    if not task["result_path"] or not os.path.exists(task["result_path"]):
        return jsonify({"error": "Arquivo de resultado nao encontrado"}), 500

    result_path = task["result_path"]
    temp_dir    = task["temp_dir"]
    filename    = os.path.basename(result_path)

    with open(result_path, "rb") as fh:
        data = fh.read()

    def _cleanup():
        import time; time.sleep(10)
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir, ignore_errors=True)
        with tasks_lock:
            tasks.pop(task_id, None)

    threading.Thread(target=_cleanup, daemon=True).start()
    return send_file(io.BytesIO(data), as_attachment=True, download_name=filename)


def _process_in_background(task_id: str, tool: str, saved_paths: list, temp_dir: str):
    try:
        files = [SavedFile(p) for p in saved_paths]
        set_progress(task_id, 8, "Iniciando processamento...")

        dispatch = {
            "pdf-to-images": lambda: pdf_to_images(files[0], temp_dir, task_id),
            "images-to-pdf": lambda: images_to_pdf(files, temp_dir, task_id),
            "merge-pdf":     lambda: merge_pdfs(files, temp_dir, task_id),
            "split-pdf":     lambda: split_pdf(files[0], temp_dir, task_id),
            "compress-pdf":  lambda: compress_pdf(files[0], temp_dir, task_id),
            "pdf-to-pdfa":   lambda: pdf_to_pdfa(files, temp_dir, task_id),
            "word-to-pdf":   lambda: word_to_pdf(files, temp_dir, task_id),
            "excel-to-pdf":  lambda: excel_to_pdf(files[0], temp_dir, task_id),
            "txt-to-pdf":    lambda: txt_to_pdf(files[0], temp_dir, task_id),
            "pdf-to-word":   lambda: pdf_to_word(files[0], temp_dir, task_id),
        }

        if tool not in dispatch:
            raise ValueError(f"Ferramenta nao suportada: {tool}")

        output_files = dispatch[tool]()
        set_progress(task_id, 95, "Preparando arquivo para download...")
        result_path = _build_result(output_files, temp_dir)

        with tasks_lock:
            tasks[task_id]["progress"]    = 100
            tasks[task_id]["status"]      = "done"
            tasks[task_id]["message"]     = "Concluido com sucesso!"
            tasks[task_id]["result_path"] = result_path

    except Exception as exc:
        with tasks_lock:
            if task_id in tasks:
                tasks[task_id]["status"]   = "error"
                tasks[task_id]["message"]  = str(exc)
                tasks[task_id]["progress"] = 0


def _build_result(output_files, temp_dir):
    if not isinstance(output_files, list):
        output_files = [output_files]
    if len(output_files) == 1:
        return output_files[0]
    zip_path = os.path.join(temp_dir, "converted_files.zip")
    with zipfile.ZipFile(zip_path, "w") as zipf:
        for fp in output_files:
            zipf.write(fp, os.path.basename(fp))
    return zip_path


def pdf_to_images(file, temp_dir, task_id=None):
    pdf_path = os.path.join(temp_dir, secure_filename(file.filename))
    file.save(pdf_path)
    doc = fitz.open(pdf_path)
    total = len(doc)
    output_files = []
    for page_num in range(total):
        if task_id:
            set_progress(task_id, 10 + int(page_num / total * 80),
                         f"Convertendo pagina {page_num + 1} de {total}...")
        page = doc.load_page(page_num)
        pix  = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        img_path = os.path.join(temp_dir, f"page_{page_num + 1}.png")
        pix.save(img_path)
        output_files.append(img_path)
    doc.close()
    return output_files


def images_to_pdf(files, temp_dir, task_id=None):
    total  = len(files)
    images = []
    for i, file in enumerate(files):
        if task_id:
            set_progress(task_id, 10 + int(i / total * 70),
                         f"Processando imagem {i + 1} de {total}...")
        img_path = os.path.join(temp_dir, secure_filename(file.filename))
        file.save(img_path)
        img = Image.open(img_path)
        if img.mode != "RGB":
            img = img.convert("RGB")
        images.append(img)
    if task_id:
        set_progress(task_id, 85, "Gerando PDF...")
    pdf_path = os.path.join(temp_dir, "images_to_pdf.pdf")
    images[0].save(pdf_path, save_all=True, append_images=images[1:])
    return [pdf_path]


def merge_pdfs(files, temp_dir, task_id=None):
    total      = len(files)
    merged_doc = fitz.open()
    for i, file in enumerate(files):
        if task_id:
            set_progress(task_id, 10 + int(i / total * 80),
                         f"Mesclando arquivo {i + 1} de {total}...")
        pdf_path = os.path.join(temp_dir, secure_filename(file.filename))
        file.save(pdf_path)
        doc = fitz.open(pdf_path)
        merged_doc.insert_pdf(doc)
        doc.close()
    if task_id:
        set_progress(task_id, 90, "Salvando PDF final...")
    output_path = os.path.join(temp_dir, "merged.pdf")
    merged_doc.save(output_path)
    merged_doc.close()
    return [output_path]


def split_pdf(file, temp_dir, task_id=None):
    pdf_path = os.path.join(temp_dir, secure_filename(file.filename))
    file.save(pdf_path)
    doc   = fitz.open(pdf_path)
    total = len(doc)
    output_files = []
    for page_num in range(total):
        if task_id:
            set_progress(task_id, 10 + int(page_num / total * 80),
                         f"Extraindo pagina {page_num + 1} de {total}...")
        new_doc = fitz.open()
        new_doc.insert_pdf(doc, from_page=page_num, to_page=page_num)
        output_path = os.path.join(temp_dir, f"page_{page_num + 1}.pdf")
        new_doc.save(output_path)
        new_doc.close()
        output_files.append(output_path)
    doc.close()
    return output_files


def compress_pdf(file, temp_dir, task_id=None):
    # Ghostscript PDFSETTINGS: /screen=72dpi  /ebook=150dpi  /printer=300dpi
    pdf_path    = os.path.join(temp_dir, secure_filename(file.filename))
    file.save(pdf_path)
    output_path = os.path.join(temp_dir, "compressed.pdf")

    if task_id:
        set_progress(task_id, 20, "Analisando PDF...")

    gs_args = [
        "gs", "-dBATCH", "-dNOPAUSE", "-dQUIET", "-dSAFER",
        "-dPDFSETTINGS=/ebook",
        "-dColorImageResolution=150",
        "-dGrayImageResolution=150",
        "-dMonoImageResolution=150",
        "-sDEVICE=pdfwrite",
        "-dCompatibilityLevel=1.4",
        f"-sOutputFile={output_path}",
        pdf_path,
    ]
    gs_bytes = [a.encode("utf-8") if isinstance(a, str) else a for a in gs_args]

    if task_id:
        set_progress(task_id, 40, "Comprimindo imagens com Ghostscript...")

    try:
        ghostscript.Ghostscript(*gs_bytes)
    except Exception as e:
        raise RuntimeError(f"Erro ao comprimir PDF: {e}") from e

    if task_id:
        set_progress(task_id, 85, "Verificando resultado...")

    orig_size = os.path.getsize(pdf_path)
    comp_size = os.path.getsize(output_path)
    if comp_size >= orig_size:
        return [pdf_path]
    return [output_path]


def pdf_to_pdfa(files, temp_dir, task_id=None):
    if not isinstance(files, list):
        files = [files]
    total        = len(files)
    output_files = []
    for i, file in enumerate(files):
        if task_id:
            set_progress(task_id, 10 + int(i / total * 80),
                         f"Convertendo {file.filename} para PDF/A ({i + 1}/{total})...")
        input_path = os.path.join(temp_dir, secure_filename(file.filename))
        file.save(input_path)
        base_name, _ = os.path.splitext(os.path.basename(input_path))
        output_path  = os.path.join(temp_dir, f"{base_name}_pdfa.pdf")
        gs_args = [
            "gs", "-dPDFA=1", "-dBATCH", "-dNOPAUSE", "-dNOOUTERSAVE",
            "-dUseCIEColor", "-sProcessColorModel=DeviceRGB", "-sDEVICE=pdfwrite",
            "-sColorConversionStrategy=UseDeviceIndependentColor",
            "-dPDFACompatibilityPolicy=1",
            f"-sOutputFile={output_path}", input_path,
        ]
        gs_args = [a.encode("utf-8") if isinstance(a, str) else a for a in gs_args]
        try:
            ghostscript.Ghostscript(*gs_args)
        except Exception as e:
            raise RuntimeError(f"Erro ao converter {file.filename} para PDF/A: {e}") from e
        output_files.append(output_path)
    return output_files


def word_to_pdf(files, temp_dir, task_id=None):
    from docx import Document

    if not isinstance(files, list):
        files = [files]

    pdf_path      = os.path.join(temp_dir, "word_to_pdf.pdf")
    c             = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    y_position    = height - 50
    total         = len(files)

    for file_idx, file in enumerate(files):
        if task_id:
            set_progress(task_id, 10 + int(file_idx / total * 80),
                         f"Convertendo {file.filename} ({file_idx + 1}/{total})...")
        docx_path = os.path.join(temp_dir, secure_filename(file.filename))
        file.save(docx_path)
        doc = Document(docx_path)

        if file_idx > 0:
            c.showPage(); y_position = height - 50
            c.setFont("Helvetica-Bold", 12)
            c.drawString(50, y_position, "=" * 60); y_position -= 20
            c.drawString(50, y_position, f"Documento: {file.filename}"); y_position -= 20
            c.drawString(50, y_position, "=" * 60); y_position -= 30
            c.setFont("Helvetica", 11)

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                chars_per_line = int((width - 100) / 6)
                words, lines, current_line = paragraph.text.split(), [], []
                for word in words:
                    if len(" ".join(current_line + [word])) <= chars_per_line:
                        current_line.append(word)
                    else:
                        if current_line: lines.append(" ".join(current_line))
                        current_line = [word]
                if current_line: lines.append(" ".join(current_line))
                for line in lines:
                    if y_position < 50: c.showPage(); y_position = height - 50
                    c.drawString(50, y_position, line); y_position -= 20

        for table in doc.tables:
            y_position -= 10
            if y_position < 100: c.showPage(); y_position = height - 50
            c.setFont("Helvetica", 9)
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if len(row_text) > 100: row_text = row_text[:97] + "..."
                if y_position < 50: c.showPage(); y_position = height - 50
                c.drawString(50, y_position, row_text); y_position -= 15
            y_position -= 10; c.setFont("Helvetica", 11)

    c.save()
    return [pdf_path]


def excel_to_pdf(file, temp_dir, task_id=None):
    xlsx_path     = os.path.join(temp_dir, secure_filename(file.filename))
    file.save(xlsx_path)
    pdf_path      = os.path.join(temp_dir, "excel_to_pdf.pdf")
    c             = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    y_position    = height - 50

    if task_id: set_progress(task_id, 20, "Lendo planilha...")

    try:
        workbook     = openpyxl.load_workbook(xlsx_path)
        total_sheets = len(workbook.sheetnames)
        for sheet_idx, sheet_name in enumerate(workbook.sheetnames):
            if task_id:
                set_progress(task_id, 20 + int(sheet_idx / total_sheets * 65),
                             "Processando aba " + sheet_name + "...")
            sheet = workbook[sheet_name]
            c.setFont("Helvetica", 10)
            c.drawString(50, y_position, f"--- Planilha: {sheet_name} ---"); y_position -= 20
            for row in sheet.iter_rows():
                row_data  = [str(cell.value) if cell.value is not None else "" for cell in row]
                line_text = " | ".join(row_data)
                max_chars = int((width - 100) / 6)
                if len(line_text) > max_chars: line_text = line_text[:max_chars] + "..."
                if y_position < 50:
                    c.showPage(); y_position = height - 50; c.setFont("Helvetica", 10)
                c.drawString(50, y_position, line_text); y_position -= 15
            y_position -= 30
            if y_position < 50 and sheet_name != workbook.sheetnames[-1]:
                c.showPage(); y_position = height - 50
    except Exception as e:
        c.drawString(50, y_position - 20, f"Erro ao ler planilha: {e}")

    c.save()
    return [pdf_path]


def txt_to_pdf(file, temp_dir, task_id=None):
    txt_path      = os.path.join(temp_dir, secure_filename(file.filename))
    file.save(txt_path)
    pdf_path      = os.path.join(temp_dir, "text_to_pdf.pdf")
    c             = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    y_position    = height - 50
    c.setFont("Helvetica", 12)

    if task_id: set_progress(task_id, 20, "Lendo arquivo de texto...")

    try:
        with open(txt_path, "r", encoding="utf-8") as fh:
            lines = fh.readlines()
        total = len(lines)
        for i, line in enumerate(lines):
            if task_id and i % 50 == 0:
                set_progress(task_id, 20 + int(i / max(total, 1) * 70),
                             f"Processando linha {i + 1} de {total}...")
            text_line      = line.strip()
            chars_per_line = int((width - 100) / 7)
            chunks = (
                [text_line[j:j + chars_per_line] for j in range(0, len(text_line), chars_per_line)]
                if text_line else [""]
            )
            for chunk in chunks:
                if y_position < 50:
                    c.showPage(); y_position = height - 50; c.setFont("Helvetica", 12)
                c.drawString(50, y_position, chunk); y_position -= 15
    except Exception as e:
        c.drawString(50, y_position - 20, f"Erro ao ler arquivo: {e}")

    c.save()
    return [pdf_path]


def pdf_to_word(file, temp_dir, task_id=None):
    pdf_path      = os.path.join(temp_dir, secure_filename(file.filename))
    file.save(pdf_path)
    docx_filename = os.path.splitext(secure_filename(file.filename))[0] + ".docx"
    docx_path     = os.path.join(temp_dir, docx_filename)

    if task_id: set_progress(task_id, 20, "Analisando PDF...")

    cv = None
    try:
        cv = Converter(pdf_path)
        if task_id: set_progress(task_id, 40, "Convertendo para Word (pode demorar)...")
        cv.convert(docx_path)
        if task_id: set_progress(task_id, 85, "Finalizando...")
    except ValueError as e:
        raise RuntimeError(f"Erro no arquivo PDF: {e}") from e
    except ConversionException as e:
        raise RuntimeError(f"Erro interno na conversao: {e}") from e
    except Exception as e:
        raise RuntimeError(f"Erro ao converter {file.filename} para Word: {e}") from e
    finally:
        if cv: cv.close()

    return [docx_path]


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)
